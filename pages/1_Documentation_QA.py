import json
import os
import re
from pathlib import Path
from typing import Any

import pdfplumber
import streamlit as st
from docx import Document
from openai import OpenAI
from pypdf import PdfReader
from rapidfuzz import fuzz
from auth import logout_user, require_auth
from research import WEB_SEARCH_TOOLS, with_web_research
from theme import (
    apply_trimera_theme as apply_shared_theme,
    page_header as shared_page_header,
    render_topbar,
    sidebar_label,
    sidebar_model,
    sidebar_reminder,
)

st.set_page_config(
    page_title="Trimera Documentation QA",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)



def apply_trimera_ui() -> None:
    """Apply the full Trimera Health branded interface."""
    st.markdown(
        """
        <style>
        :root {
            --trimera-green: #76c64f;
            --trimera-green-soft: #eef8e9;
            --trimera-blue: #168fd0;
            --trimera-blue-soft: #eaf6fd;
            --trimera-teal: #0b6b73;
            --trimera-navy: #102033;
            --trimera-text: #13283c;
            --trimera-muted: #6b7f91;
            --trimera-border: #d9e4ec;
            --trimera-bg: #f7fafc;
            --trimera-card: #ffffff;
        }

        html, body, [class*="css"] {
            font-family: Inter, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
        }

        /* Full app background */
        [data-testid="stAppViewContainer"] {
            background:
                radial-gradient(circle at 92% 6%, rgba(22,143,208,.08), transparent 30rem),
                linear-gradient(180deg, #fbfdfe 0%, var(--trimera-bg) 100%);
            color: var(--trimera-text);
        }

        /* Fixed branded top bar */
        .trimera-topbar {
            position: fixed;
            top: 0;
            left: 330px;
            right: 0;
            z-index: 1000;
            height: 74px;
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 0 28px 0 30px;
            background: linear-gradient(
                90deg,
                var(--trimera-green) 0%,
                var(--trimera-green) 48%,
                var(--trimera-blue) 52%,
                var(--trimera-blue) 100%
            );
            box-shadow: 0 6px 20px rgba(16,32,51,.10);
        }

        .trimera-topbar-brand {
            display: flex;
            align-items: center;
            gap: 10px;
            color: white;
            font-size: 1rem;
            font-weight: 760;
            letter-spacing: .04em;
            text-transform: uppercase;
        }

        .trimera-topbar-brand img {
            height: 28px;
            width: auto;
            object-fit: contain;
            filter: brightness(0) invert(1);
        }

        .trimera-topbar-app {
            color: white;
            font-size: .92rem;
            font-weight: 750;
            letter-spacing: .06em;
            text-transform: uppercase;
        }

        [data-testid="stHeader"] {
            background: transparent;
            height: 74px;
        }

        [data-testid="stMainBlockContainer"] {
            max-width: 1180px;
            padding-top: 6.4rem;
            padding-bottom: 4rem;
        }

        /* Sidebar */
        [data-testid="stSidebar"] {
            width: 330px !important;
            min-width: 330px !important;
            max-width: 330px !important;
            background: #ffffff;
            border-right: 1px solid #e4ebf0;
            padding-top: 82px;
            overflow: visible !important;
        }

        [data-testid="stSidebar"] > div:first-child {
            width: 330px !important;
            min-width: 330px !important;
            max-width: 330px !important;
        }

        [data-testid="stSidebarContent"] {
            width: 330px !important;
            min-width: 330px !important;
            max-width: 330px !important;
            padding-left: 14px;
            padding-right: 14px;
        }

        [data-testid="stSidebar"]::before {
            content: "";
            position: absolute;
            top: 74px;
            bottom: 0;
            left: 0;
            width: 10px;
            background: linear-gradient(
                180deg,
                var(--trimera-green) 0%,
                var(--trimera-green) 48%,
                var(--trimera-blue) 52%,
                var(--trimera-blue) 100%
            );
        }

        [data-testid="stSidebar"] * {
            color: var(--trimera-text) !important;
        }

        [data-testid="stSidebarNav"] {
            padding-top: .6rem;
        }

        [data-testid="stSidebarNav"]::before {
            content: "APPLICATIONS";
            display: block;
            margin: 0 1.2rem .5rem 1.2rem;
            color: #657789;
            font-size: .72rem;
            font-weight: 800;
            letter-spacing: .08em;
        }

        [data-testid="stSidebarNav"] a {
            margin: .2rem .45rem;
            padding: .72rem .82rem;
            border-radius: 12px;
            font-weight: 620;
            transition: background-color .15s ease, transform .15s ease;
            white-space: normal !important;
            overflow: visible !important;
            text-overflow: clip !important;
            line-height: 1.25;
        }

        [data-testid="stSidebarNav"] a span,
        [data-testid="stSidebarNav"] a p {
            white-space: normal !important;
            overflow: visible !important;
            text-overflow: clip !important;
            max-width: none !important;
        }

        [data-testid="stSidebarNav"] a:hover {
            background: #f0f6f2;
            transform: translateX(2px);
        }

        [data-testid="stSidebarNav"] a[aria-current="page"] {
            background: linear-gradient(90deg, #edf8e9 0%, #f2f9f4 100%);
            box-shadow: inset 4px 0 0 var(--trimera-green);
        }

        [data-testid="stSidebar"] [data-testid="stButton"] button {
            background: #ffffff !important;
            color: var(--trimera-text) !important;
            border: 1px solid #dbe6ec !important;
            box-shadow: none !important;
        }

        [data-testid="stSidebar"] [data-testid="stButton"] button:hover {
            border-color: var(--trimera-green);
            background: #f5fbf2;
        }

        [data-testid="stSidebar"] [data-testid="stAlert"] {
            background: #ffffff;
            border: 1px solid #dbe6ec;
            box-shadow: 0 8px 20px rgba(16,32,51,.05);
        }

        [data-testid="stSidebar"] code {
            color: #0b6b73 !important;
            background: #eef8f3;
            border-radius: 6px;
            padding: 2px 6px;
        }


        .trimera-sidebar-brand {
            display: flex;
            align-items: center;
            gap: 12px;
            margin: .1rem .25rem 1rem .25rem;
            padding: .75rem .8rem .95rem .8rem;
            border-bottom: 1px solid #e7edf1;
        }

        .trimera-sidebar-brand img {
            display: block;
            width: 74px;
            height: auto;
            object-fit: contain;
        }

        .trimera-sidebar-brand-text {
            display: flex;
            flex-direction: column;
            min-width: 0;
        }

        .trimera-sidebar-brand-name {
            color: var(--trimera-text) !important;
            font-size: 1.22rem;
            font-weight: 780;
            line-height: 1.05;
            letter-spacing: -.025em;
        }

        .trimera-sidebar-brand-subtitle {
            margin-top: .28rem;
            color: #6a7d8e !important;
            font-size: .72rem;
            font-weight: 700;
            letter-spacing: .06em;
            text-transform: uppercase;
        }

        /* Page header card */
        .trimera-page-card {
            position: relative;
            overflow: hidden;
            margin-bottom: 1.4rem;
            padding: 1.45rem 1.6rem 1.35rem 1.9rem;
            border: 1px solid #e3eaef;
            border-radius: 18px;
            background: rgba(255,255,255,.96);
            box-shadow: 0 12px 28px rgba(16,32,51,.07);
        }

        .trimera-page-card::before {
            content: "";
            position: absolute;
            inset: 0 auto 0 0;
            width: 8px;
            background: linear-gradient(
                180deg,
                var(--trimera-green) 0%,
                var(--trimera-green) 48%,
                var(--trimera-blue) 52%,
                var(--trimera-blue) 100%
            );
        }

        .trimera-kicker {
            margin-left: .3rem;
            color: #0b6b73;
            font-size: .74rem;
            font-weight: 800;
            letter-spacing: .1em;
            text-transform: uppercase;
        }

        .trimera-title {
            margin: .35rem 0 0 .3rem;
            color: var(--trimera-text);
            font-size: clamp(2rem, 3vw, 2.7rem);
            line-height: 1.08;
            font-weight: 790;
            letter-spacing: -.045em;
        }

        .trimera-subtitle {
            margin: .5rem 0 0 .3rem;
            color: var(--trimera-muted);
            font-size: 1rem;
        }

        /* Labels */
        label, p, [data-testid="stMarkdownContainer"] {
            color: var(--trimera-text);
        }

        [data-testid="stCaptionContainer"] {
            color: var(--trimera-muted) !important;
        }

        /* Select box */
        div[data-baseweb="select"] > div {
            min-height: 46px;
            border: 1px solid var(--trimera-border) !important;
            border-radius: 12px !important;
            background: #ffffff !important;
            box-shadow: 0 1px 2px rgba(16,32,51,.03);
        }

        div[data-baseweb="select"] * {
            color: var(--trimera-text) !important;
        }

        /* Intended billing input: dark background, clearly visible text */
        .stTextInput input {
            min-height: 48px;
            border: 2px solid transparent !important;
            border-radius: 12px !important;
            background:
                linear-gradient(#182234, #182234) padding-box,
                linear-gradient(
                    90deg,
                    var(--trimera-green) 0%,
                    var(--trimera-green) 48%,
                    var(--trimera-blue) 52%,
                    var(--trimera-blue) 100%
                ) border-box !important;
            color: #ffffff !important;
            caret-color: #ffffff !important;
            font-size: 1rem;
            font-weight: 600;
        }

        .stTextInput input::placeholder {
            color: rgba(255,255,255,.62) !important;
            opacity: 1;
        }

        .stTextInput input:focus {
            box-shadow: 0 0 0 4px rgba(118,198,79,.13) !important;
            outline: none !important;
        }

        /* Note textarea */
        .stTextArea textarea {
            border: 1px solid var(--trimera-border) !important;
            border-radius: 12px !important;
            background: #ffffff !important;
            color: var(--trimera-text) !important;
            caret-color: var(--trimera-text) !important;
            box-shadow: 0 1px 2px rgba(16,32,51,.03);
        }

        .stTextArea textarea::placeholder {
            color: #98a8b6 !important;
            opacity: 1;
        }

        .stTextArea textarea:focus {
            border-color: var(--trimera-blue) !important;
            box-shadow: 0 0 0 4px rgba(22,143,208,.10) !important;
        }

        /* Radio buttons */
        [data-testid="stRadio"] label {
            font-weight: 600;
        }

        [data-testid="stRadio"] input:checked + div {
            color: var(--trimera-green) !important;
        }

        /* File uploader */
        [data-testid="stFileUploader"] {
            padding: .8rem;
            border: 1.5px dashed #62bce7;
            border-radius: 16px;
            background: rgba(255,255,255,.86);
        }

        [data-testid="stFileUploaderDropzone"] {
            border-radius: 12px;
            background: #182234;
        }

        [data-testid="stFileUploaderDropzone"] * {
            color: #ffffff !important;
        }

        /* Success/alert boxes */
        [data-testid="stAlert"] {
            border-radius: 12px;
            border-width: 1px;
        }

        /* Primary buttons */
        .stButton > button[kind="primary"],
        .stDownloadButton > button {
            min-height: 48px;
            border: 0 !important;
            border-radius: 12px;
            background: linear-gradient(
                90deg,
                var(--trimera-green) 0%,
                var(--trimera-green) 48%,
                var(--trimera-blue) 52%,
                var(--trimera-blue) 100%
            ) !important;
            color: #ffffff !important;
            font-weight: 760;
            font-size: 1rem;
            box-shadow: 0 10px 22px rgba(22,143,208,.16);
            transition: transform .15s ease, filter .15s ease, box-shadow .15s ease;
        }

        .stButton > button[kind="primary"]:hover,
        .stDownloadButton > button:hover {
            color: #ffffff !important;
            transform: translateY(-1px);
            filter: brightness(1.03);
            box-shadow: 0 13px 26px rgba(22,143,208,.20);
        }

        .stButton > button[kind="primary"]:disabled {
            color: rgba(255,255,255,.78) !important;
            opacity: .58;
        }

        /* Standard buttons, including Sign in */
        .stButton > button:not([kind="primary"]) {
            min-height: 43px;
            border-radius: 11px;
            border: 1px solid #0b6b73;
            background: #0b6b73 !important;
            color: #ffffff !important;
            font-weight: 700;
            box-shadow: 0 6px 14px rgba(11,107,115,.14);
        }

        .stButton > button:not([kind="primary"]):hover {
            border-color: #0b6b73;
            background: #0b6b73 !important;
            color: #ffffff !important;
            transform: none;
            filter: none;
        }

        [data-testid="stExpander"] {
            overflow: hidden;
            border: 1px solid var(--trimera-border);
            border-radius: 14px;
            background: #ffffff;
        }

        [data-testid="stChatMessage"] {
            border: 1px solid var(--trimera-border);
            border-radius: 16px;
            background: rgba(255,255,255,.92);
            box-shadow: 0 6px 18px rgba(16,32,51,.04);
        }

        hr {
            border-color: #e3eaef;
            margin: 2rem 0;
        }

        @media (max-width: 900px) {
            .trimera-topbar {
                left: 0;
                padding: 0 16px;
            }

            .trimera-topbar-brand {
                font-size: 1.12rem;
            }

            .trimera-topbar-brand img {
                height: 42px;
            }

            .trimera-topbar-app {
                display: none;
            }

            [data-testid="stMainBlockContainer"] {
                padding-left: 1rem;
                padding-right: 1rem;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    st.markdown(
        f"""
        <div class="trimera-topbar">
            <div class="trimera-topbar-brand">
                <img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAPwAAADYCAYAAADPlo5HAAAAAXNSR0IArs4c6QAAAARnQU1BAACxjwv8YQUAAAAJcEhZcwAAEnQAABJ0Ad5mH3gAAD79SURBVHhe7Z0Hc1tJdu+bJAIDQIqURhppws7szM7suOyq9xHfF3S5bJf9vN6doCzmhEQCJN/5/c9t4gICM0gi9B+6AgHc0OHkPt09c2YICQkJU4HZ7D0hIWEKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBg+IWGKkBh+ipGmUUwf0uSZKcSpvTon7XBydhLmZgp2zIW52bns14RJRmL4KQPd3Tlrh/3WXqi1a2F+bj5USst2VLIzEiYZyaSfMpzZq33aCVvNzfB6/9fwofY2HB7vZ78mTDoSw08RYPbjk6Ow29wOO82tsN3YCpv19bDX2g3NdiOcnJ5kZyZMKhLDTwkw5U9Ms7c6zbDV6DL5wfGBBMBua0c+fcJkIzH8FKF90g7148Pw8fB9ODjaU/AOjb9nzL5e/xiOOq3szIRJRWL4KYC0u2nvA/PV182Er7froWPaXr/Zq3nSNA2/LcZvdhr6PmEykRh+CnB6dhqO8N2NobcaG9LqfIcg4DjqHJnGP1Ag79DeEQJ8nzB5SAw/BcB0R3PvNmH4TUXpI7q+fSN8qn00LW+m/ulp9mvCpCEx/ISDyDsmPH47JvuJvWay3yLQ6Jj8DTtvt7Vlmn7DhEI7+zVhkpAYfkIhc91ex6fHoWZm+mZjPdSOD+17M+Xt1Q+0+vFJS9F7zj2yvzk3YbKQGH5SYWocPx2tvX+0q6AcQ3IXws7vnHXMl98P67VPuq6ThukmDonhJxRoeIbhGGP/VPsQjjtH55p9xl6zM7N2zIUZe+ezLAKuMVO+3j4MG/VPYd+0fcJkITH8BALN3skYl4y6HWP6fFINDF6YnQ/zhWoozS2E2dlC9gvXnsgS2Kivm2WwZ0LjeKALkDCeSAw/gcD3bnVaisjvmTnfOmlKCACYfcYYfKG4ElbmX4ZK8VkozS7qN8B5aPn9ox0N4xHwI4qfMBlIDD+BYNitflyTKX94fCBTHciUN2Yvzy2GxeKTsFR6GirlZ6FcqJh5j5b3+D1Mz9g8UX25A6blEyYDieEnCDA2zEqOPNp5h0Bduxuom5mZCXMzxbBQWDFzftmOiml6e7ejZEIAvx5wH4bzDlr78uWJ7mPaJ4w/EsNPErJIOxH5zcYn5cafZKY8MP0eisx/N60OsxOwm5stmQComrZfy7S8g2QdhuZwCbabm5o7nzD+SAw/IVCE/aQdGkc11+4K1OF7uzmP9kaLo93LhSVj9KKb+DNzoVRwE5/fWQEngkAfQoMAHuPzWA/RPUgYTySGnxDEmW+7ppFhzrpp5Di/Hc0+G8x3N62+WFwNxdl5Mbr/RsS+rN8w7wtmAUTExB18eQ4CeDwnYXyRGH5C4Cm0tfDu4LWG0/KaGN991jQ6fvtiafWc2SMQBzL1S2bqz1UlBIDHBHyYjvH89dqHNIV2zJEYfszhTOkZdWh2/PcjY9DzJBv56eWwZJqdcfeC+ex8l4cH87AAlsJ8kbH5pXOhwP2ZSkvgjriAMvDSMN3YIjH82MNnu7Eu3U5rS1o+P/EFxkV7E5SDoWMkvhem0+17zkMoLJg/nw/gaWJNp6G4AELl0hTdhJFGYvixhmv3tvnZJNms1z6e++0A351hOMbdl8yUJ6vuIiiAZ34+DF8tP5VfH017gFBpdOoalyd7L2E8kRh+jHFq5rYy6pqbCqox/p7Pl0drw8BLxbWBpnw/MO0L5uuXzaRfLK4oah9BsI5RACbiMFRXax8m034MkRh+jOGmdl2anYw6xuBjsA7mhskZhvNAXddEvwycJxfAfH4i91HLc1/G9HmeYgVm3uM6xOcljAcSw48tGHc/FqNjZpNKe87s9sKUJ+JOEA6T/irtHqEAnmn5hdITsw4q2Xi9X0uOPloehn9/+NasiwYPSxgjJIYfQ2C2Y06TYLNZ3/AJLqbtI2ayQJ3nyS9dm9kduAJzsg4YxsMdmMvNplMyzkkz7DV3w05jW4ImYXyQGH7MALMTQGNRSgJ1bCTBVFiCdwDtjh/O0BqBuuIlgbqL4BYCAbxKqJTw/8vnQgMr4rhzLB+e2MH+0b40f4wdJIw2EsOPIeK4OKY1Jn2e2dHO+N5LZpLD7GwUeRtIyzNMVyQVtxIKM6Xsl25W3zbTb1vbWicvYTyQGH7MEDPqPhy+M2bfk3Y/h2lhfG5M8e5Y+u2d7PMx/CyAd34vU+YewPNZebgVCICE0Udi+DEBpjQHY+5o9Y+1d74o5bkpPSMGZygNhi8Xqvb5bt3L9fjyTKxZsHsSCCSAxzMx40mzJY2XsjTajXNLI2F0kRh+jEBeO4tM7jS2fBguNw7us+HmQ7X0hXzvuzK7w12EUmFB0X7G5vMBvFMz5QnabTY2TPhQnrS09agjMfyY4MzYi33dd8xvZvcYhse6GtWH0kiY8UDdPKya/XY3cJ9ZsvUUFyBiXzq/NxaH1s7LmJ6YQsJoIzH8WMAZq3FcD9utLfnNeWYngk6ADjOeAzN8mOD+pOUulnzOfH7RS2IKrXYrbNY+KYiH1ZEi9qOLxPBjAAJkNWN2kl1Ydio/OQZtSyQeP7tSemra+HZR+cvgWr6gxS7R8qTrRsDaROnrZOAd7Wrxy7Qc1ugiMfwYAPOdAN3H2gctNZUfhsOnLs9VFFTzJJv7SX1Dy+M24MfD8B7A41kewCNKT3yBLa2YTRez/hJGC4nhRxgwDcxNCiv7uW83t5RwE+GTXcphwfx2jZXn/OthQ1o+Ey4wvEz7nDWBKc8U3Y+HH8Lh0WFP5l/C6CAx/AiDQB3anVz5j7X38uPRphGeQrsQlsvP5GPfN6JFwbBfdf65/V3OfmEEwdezJwOPvelYiCNh9JAYfkRB4Iv15RudmpabZtun/kAdUXk2lEDrYmI/BFzIzMu0x4XIBwgJ4LHSLeVla2oW0cwLqITHR2L4EQXmPIktMA6BumZulRlMeYJorCdPFpxM+aGMu18NtDzPm1fcoBqKs13LImp5nzO/ozIn0360kBh+RBF3cv1979e+jDo6jQw4893ZUMKY3rg9++VhQLYd4/GM+S8UqxJAEXlB9Wb/D2XgJYwOEsOPGGAYzOD60aHG27UopZnJEWhyJrX4GnX3G6i7CG5hzIUSATwTOIz959ezJ4DHYpdx15qUgTc6SAw/YkCTwzC+mcSWlq3Km8XOaIthufzcfPfFB2f2CBc8ZQXwKsWn0vgRWnqLXWusDowuJC0/OkgMP2LAD2bzB9Jn2eIpDzQrm0jgP2ssPMdkjwHy9RklWCyuyZefCXGYzvemY729T7WPmrOfgnejgcTwIwTmmaPRYRA0PJox+u7ngTrTqEx9xZQfzgSZ2wNfnnny84VstKDQXfRSuf9myh8c7ylqf2imfcrAe3wkhh8RwNhoxYPjfQW78H2ZHReB9pwz5iKfncUtHpvZI2Zn50Kh4DEFmD66GDFpqNlpKGq/1dzoSRpKeBwkhh8RxOg200wx5/NbOkmTzpYVFfeFJTHlH8d37wcMbqJIY/IEERmjz2fgKS346DB8OHjne9Ml0/5RkRh+BIB/yzp1RORZp46loPPLRnmgbkHBsf6U1lGAAngmhOIa+PnYAoKsxaKXjDg0tz4bYkx4WCSGHwEwG47JJ0w8YY15tCCMEkEOe8k0KKvQErQbRcS96ZbnX2RlzEx7ezHqwLJcG41PYae5mbT8IyIx/CMD4se33W3tyn9HG0ZgLhOoY4tnMuq0RvwDJ9lcF5SLAB6r5eLLI6AiEF6eSHSgehK8S0z/OEgM/4hA+0H4mvpaf69Idv/68pjK7NvO4YG60WR4yoWrgXaH4Yk1xLLGgKQH8PbCRn1dU2gTHh6J4R8JaL2zU989hokxr3d/C812/dyUd+3uDIRvzESZ0WV2B758nDPP/HwEVD4x6KhzJF/+zf7vEm7Jl394JIZ/JEjr2Ys57j7PvSUtGAHzwOSrC1/pfVRN+X5QTIJ2CKlq+bmy8SKwZnyhjL1waEer3d3HPuFhkBj+ESDtnhE/2XRErzHlI/GjFRmGmy/6wpFFMc14MDzljAE8ltzyDLxIZlg0ba1nv0PUvrUtIZCY/uGQGP6RAIPjx+60NrUWXATMjjZnGE6mvDEOgbtxglJuzRUh2Eg98uX3IGUrbNY/aaGMfHJRwv0jMfwjAGaHyf/Y+1WR615THg1ZzDLXVvnGfxgrILTIs2cK7ZqyA/M4PSUDrxnYly4F8B4WieEfEDLlMWvJMW/ta9mqZs6PRbvD7GSs9eemjxvQ8tRFIwyFZQ3ZYdp7C5yZO3Os+f5MoWUqbZpc8zBIDP/AgLCPTKOxcwxDVGwdFeFacSFUy18oUPdQy1bdD9DyniHIEB0Zgvk582QWsoEF6/Vp3oC98PET7heJ4R8QEHWrcxTeHr4J63XPqOvCh+Fg9ErpmYJ2k4C4a83ywpcSZtQzgjXvfKGMdR3Mo0+4XySGfyDIlGciSftAwSq2ZfJoPUQOs/swnK8gU+nZ3WWcoQAei17mltIGtAexi7joJYt9YObHPISE+0Fi+AeAvFYj5GanLsJm9lh+e2UP1BWk2Zl8gqbHn58MIMxYzz7Oma+c14120Xr25ssjAGF+UnAT7g+J4R8CprTIqIOomevOIhd5wBCsUwczlMY4UHcRlIFndWTBSwJ4+PZ50/4823D/NwXyEAQJ94PE8PcMma5nJwpM7TS3w3YD07Wr3VmBtmj+OmPWvs77OAfqBgONDtMzsSa6LPltp9HqzKZjpiATiHB9EtPfDxLD3zMw5fFNmRpKVh1R+Twxo+1IsFmb/0qR7LzmmyTA9Agz1rJfKb8Ipdx69t5GbeXXo+Fh/t6AZsKwkBj+HuHavWO+aVNbKUPMXWRRedPqMDzTSfk8yWBtHK2nb748EfveDLwTrfJDOzE23zlh15qk5YeNxPD3BJidF/46k2MYc8/77gTqYHBflHI5m+s+2d2h4KTVk+W1JeRyy2zTVix6iS+PJYQLlN8WO2E4SAx/j0BLwey/7vyvtnmGqCOM3TWrLC7+OC1AqDGDjhEJJtfkhRztw2Ig+6095Sk0rM0ShovE8PcE8sXJET88OjBTfs+Yv6utYPa4Rh3aDtMW83UaDvtP9WcGIO4Mgbx8Bh7tRDIOq9zWjpMvP2zMWCckR2nIgEgJ1H2qf9BQ05u93/W9kbz9z5j7nLR6tfwiLJo5H5NRpgmk1rY6BOnWzdU5MPM9rtJLcI9lvZbCj2s/h29XvjfBWDGhmHTTMJAYfsigOUkmwQf9z41/U3QejdUFfiwLRFSMqJ/Yp+kFgvHkzDT68Z4JyHomEM3sxOw3X//50pfhq+o34c+rP4biFArF+0Bi+CECgqU5YXZ2j/mf7f/SDjJxzjemLCmz+K+Y9KDZ3jdtN30bNMwaQxOpZwfcVudAWv7IND5tSHCPtlouL4cXS6/CL8/+OSyV2AN/skcxHgKJ4YcIiBXtzuIOr/d/D+8O3mjhxgh8VYj86eJ3WiDiyLTaXvO9TFu2ZpoesMDHojHxalhb+NraoRHqxzthv/XJhKOn1sLy5cJ8WJ1/Gn56+tfwbPG5zPwY1U+4HZJjNGTEnV/fH77tyaiTzpopitBlzpeehGrpqQJX+ayzSYe3A0FLX5yTA9eGhBw0eGRohGdc/46FQthvPummuyMx/JAQ16gjaYTkESL0PsfboYy6YiWszH8pcx5tH/d5x6ydFs3lS2+TSrzmE4VM2PG5XFwOlfJzs4ByG1Iag9Omce0A5s/nl/FOuDkSww8J7B7DKqyMH7NAIwGpqJE09jxb0sQRFreAwOMMOWk3Oxiaw2+dZLh2L5jAW7J6r5gArKrOWDgMT1ZLLPxBMo63A1oeBmc5LCYesYkFU2r5PuF2SAw/JJAlhr/O9Fc0UYQTuefLcxCh92QT+8WYHp9es+SMCSY+tdbqzfg7Qs5Ta72+zvQluTd8TzYe3wKYm5Rbtqhar3/QzLqE2yMx/B2BFne/fTu83v9DzH6a2wgSImebZ0x30mh7TXfX8vj0T+ZfytSfVNOeemHZIPSqpWfG2L175GnWoAkDAnlsqxWFAe2LtYSLhFm/2dzoEagJN0Ni+DsCDdQ2P3O7uR3eHb4RYXaDSzA0Szxhwj7Rez/Q8jC61rEzxp+UlW7ygNldu7O+3XJYKD2Ri5NHFIwkJCEc+T2fbMNsOrbSZg08shdp4247J1wXieHvgOhjsu95/fhQk2PygToIFo22Us60d46Au3CTHwIv27msDDNpEDObmb5Y8C2oEHKDwPe0AzMHOW/OLIII2rllpv2uCdZa+1AulDVdwg2RGP4OkKnZbkqzs687Y+lR66DVFKDKtDs+6oUwQkezY86yEAZm/qSY9rEdyrOL54tfXFY3D+othqUye+HnF72MGYw1a+sNBUfTQhk3R2L4WwJCO+4wTrwfPhy86xsnxoRlieZFETjLVl0WkHOm8KmyHtDyCTWTAJnypqkx1YlV+LZZF4N28MBeXPSyqPYBWuL7pBW2mpuap8DfCTdDYvhbAuJTCm3jk6Lz+Uw5iJwA1ZP5V5oCet2JH0pIMeGwzNCdGGP8tTwxDLk183GZ6qthDoCdy1ZVT7TabXSFELJxoQwCd7hQaPmE6yMx/C2AJse8ZOvjj7UP4bhn2SqPvHcXeeiNRl8G14Yk4zBttnfa6LghWi1x2JHj2vUxFyeOXpCUpMVBcsKPZcJkWR2+04IZyay/PhLD3xAwuxalbNe0nrrWqcuNDfvY+ryZsMti9puY5hC1EnTM/CXV1H3Y8QTtQFwCpu3NP7gaUVjENGSN2edGL1hYBF/+w+FbpTEjfFPE/npIDH9DEKjDhH+z/7tWs+FzhBOqZ89hzuf3Rr82YBQjdrQ8w1P2KfthfBDbAWulUn762WaS1wX3YOSCtkRocGcg0z6chNZJ0zT9rlla2yaE03r210Fi+BuChRsax3XlyzMe3IUzKoRJ8I3ovK+/fjPALDA5vjzE7gG/8TLtYwyDlGGtaHPZCMUl8Pt4SrKvjNPNwGNFIRYZYRdetqk6sr+TaX81EsPfABqGO2nJfyTrKx8ldhO2KO0O08P8sO5t4OPRvgQU0WqIfnzg2h1zvFIkgaZ8+3awF7684gDG8Lg4MQAKc8PwrH/HkGjDTPwUwLsaieGvCQgMX5FMr3/s/k1JIHmglWFMT6GtZN/eHmg3Al2Ys7wjQEYfxqKZsMLKwS3h7ztB92O//NVQLT/PfPl8AI/17PfD+9pbCeGk5S9HYvhrAmZngwRMSAgr7zM6c7Lrq2/zPIz02KjduC+zynxI63aa8qEAs5/vCZ+t1XfdIcmLQDt4AG/B7ukTkPK789AvpDMrA+/4IM2muwKJ4a8Bxn47p8dGVFsiKqLyvRl1cxo+Wpl/oUAd3w0DCBKGpPBhYSA+jzLTM36ORset8fIOq6wmSKwdcBOYVpsfz8fNwrTHzeJgKi3+fcJgJIa/AmgLhuEOjw7D73u/Zim0XQ0ippwh/5vx5uFOcT0XJkbkLBhB1PuuGvO+QDug0T1oiRa+oynfBwKgDHfiMsH4LvwcaHVGTki3fW19RPQ+YTASw18BNDljvoy3K1CX893x2zFhK+VnpoXjIhbD1cDcz31iVrklgDdcRhoO3P1gZIIFOjVuPkTBB1z4+eIZnqew3PMMTHu24WZcvkEGXtq1ZiASw18CJdmY9iCbi6WrWG6pX7tD3KSNMtY8bGaP8LniPoVW2m2kuo1aw/AMn1WVFkzu/H1A7T3jLgMBwThcSZ/ItDe3i5Rb1sHrnaacEJEY/hKQ3sE0TLQG68ujRSLwTzExGSt3U/v+oujRXMYv9swzGGo0fHnagSBluVhVWygNdmi+ey8kWmbNtLf2Ls8xsabcp+Xb4bB9EN4evFaeRArefY7E8BcA7YD5/qn2XstWMQzXzaqTMS/mq5YZeiqJGO8TELan7Jo5a8eo+PK0g/a3N5cG39196/trC9rZYwVL5/58RAzgYZExz+HwaL9HSCckhr8QDLuRL0+gDobPm4cafjJNhmm5XHquv+8bEDpML3OW1V5nHn+3WcqkocNZY77iqhj+IaBFL43h2Wfelw3rtgOzGDHnWfSSabT5EZWExPADgSlI0IctolhdhR1NI6RhzF8lOCWNZibmfWv3iKjdPANvReV4TMxa3QnULZXX5Gbclyn/OawlTPixzDcanvhGFH5oeTLuDo73w2ZjXZmR+anL047E8APAMNxOczt8NHOeiG/XlDdSg9DMhEXL4kc+FLNHnKetaruq3uGph4THFUj/ZfRg7UGsnDzQ6nEY0FcJyuXZ24vxeBYl8b3ma/o+ITH8ZyAqjxn4sfYu/Lb7d0Xmu8Cs9txuovKMvT80YLQYse/OMX9YoYOQ85gC01c9KejBGR63ypgcS6dafibmz1sY9CELlLw/eKMVhUEK4iWG7wG+HkM7aHe0AszfhRM5Od2rC69EYA/NaBEwHAzm49GM/z9sN8rKMXcCK2e4GXU3Q4ylsNgI/cJ7BH2JW6YFRu0gAJt8+cTwPZApaCY8pjxjub1j7k5cbkJ6ks2jwcrC8ykLTO/a7WG6UnEE06xYFzD7Yy/SEUcv6JN8jgJ9h2tGHIag64b586yUM+1IDJ+DsrWOD8Lrvd807TJCJmyYC6VZpmkuiMAey3cGblL7bDpcC0xrN+3vH+5SzJtlsWJtsZRZOo8HysN8eywdsvCwPmgfwTQ6EXuCd3/s/+qbhOTiMdOIxPCGqA22Gxvhg2n3/rRMiApNtrr4dVgw0xGWGwWg3SjXcvl5Ruz3Wy6JPTPlSbCpzrNH3mjM06d/mLTEgpfL5RcSAIB+hcG1DZhZbmxKiQCYZiSGN8jfOzHfvbWtvd3R9F1zPg6FEY02LWJaflQAgxMpp1wM1cGM59ptyHCrYs6e44tz0h7DmAY8DHjZfP28pWxXm7wFxjAdpj3r2TMfYpq1fGJ4A0k2vqvJjtJomQ4bAVMRFcdffUhf+TqA0H3OPKu7smvNIgXOfh0u3HQu2HOWNT+f596XcLkNcHHoH5YEKyjVuSuM3FU7DO8P3ykgC8Pn4zPThKlneJibqa9/3/1/58M3EZGh0KDV0vMeIhoZKIDHFFomlKxZeXuXdB4GFLC0++I2LNozyD8YRdA/xBWeMJmpuKJyR8DgbBxCyi270OZnPU4Tpp7h8ddrbd+kEC3QhTGSovLLHqAykxnGGjXA3IoxmGYjYu8r7gx3TBzfPebxY00U5kbDd+9H1PKUE7eDvenyFhlDrpj06+a2kYE3jab9VDM8Up8gTu3oQCvR9s+Gwz+ukq9tJuwomfL9cEskWxGGqL2ZtHw7DEQrB0HCLjpya4Z07+HDhR8BPKwRCb/gQpo4DcE7Nv1kJh1BPAK104apZXg6m5lV+HV/7P/ew+wQDuYhG0n4kkrX3z3msQChU07P8a/I374r08PYPs7tG0ESJ3io4b/bIgo/yktbYI3kBZQsuiwDb6u+IUEwTQk5U8nwaHYYnOmTdDqBnK60R0vMSDuQo42pDAGNOs4J3Zhd6+Ibk941Aw8h4hl9zMOvSnOOsqXjoP98GFWjF3MIv64LQvYkVt12Y0sz6hD6ieEnHHQwQRtWsam3D3uj8nag1dgxRbvHjAGzn4PgWsagi4XVc3P2tnDz2Pe6w63x1hl9uPDDQltQIJO4Q0QU9o0OKbc1rYV3Yq9pwVQyPGYdSRjs635g73nA7GhHjoed8nl3QOhkmmn6bGlFU1cRALeBuwgLHrC0thgHKycPCSvT7NqTXtZOLgPP3hiGZSYdW4YxRj8tAbypYng0Ox3LYofsC0eOdf+ilPmprz4MNz4MDyBqCF1j88QfFMC7GVxDMimFRBbmut/vEl73AerAfH2YHabHJYl1gA60HJYJe9Ju2TKsP7tyUjFdDG8v9oZjT/cPB2/7ZsNZY8z68Nbawtcyi8cVxq7KCKwyZ960/E397qgdSbIhMj9u2j2iK8BXw8r8S/0dtbyb9m2Z9PjyRO/5btIxVQwPg7Mwwl5rT9L9JGfGYbrHJZCJ7D70lNNhgrpoYUljdqXBWr2uy/Qay54pKeClaL9ZOTcVGKMEyk7Qzl20/gy8E2P0uhYoVeDWlMGkM/3UMDymPJs/ej71bjaJwjsXqY9WhDmIzLNeHN+OM6SljcB915oVZ9xr1AmGgDE8P78y1swO1LOzxfPEJJ/O6+1AsPb49EiLXuLeHZqW7x2enTxMBcPjs6HdMd3+tv3fSq/skeTSAqyesioTdtz81UGA0KkH/it+uMzZK5gXy0DaUP7/asYc4w/aAQ3PTDpm1Hk7dIUfaxbuHe2Fj4fvlJAzyZgOhrcXwzAsasGilO2cFIcxCGwtl7/sCexMAqibFqtQTkHMwBsMFxBk1FUUtFRG3Zhr9wjVzV5MqlEGngm0fD/jyzfaNa1yC51McgBv4hmeZYsx3dDqTH896hzpcwTDWHGNOMy+SQOELRPdNJsH8CD0z017mBuBgN9OdqFn6k0OsF7YdRZmXzLhh0CjHeLITbPdVJ49MyZZKGNSMfEML9/dmPzj4XulU+aZnQ6Xf8d4s5m+TgSTBRg5RqqVLSfN3cvw0n9od2MGjVubYBjnoOVg0NdM762Y2+bbYeXryBRpIvUfzKwnzoMg4Jg0TDTDS3KftBSFRXrnTbXo4+LfLs+/uHZQaxwB0/syUObPm3brzy+IVgCmPBq+38edHNDnLuCxeHB1IqAV4jxk37GmPa7fJAbwJpzhvQNZtoqVS+nUCNd8DD+tSvs5kU8mEGQINDQ49c1rN7Q7wgDi90SdyQjUXQQNO86Vra7ZKIRYwIUbi5gyesM+88yoY1QnTzOTgImkcoJ0vJj/fNDaMzPtjYIxEc4AvgIt0h6zflK1ewQCjdl/brL7hJJo5bBslXx8tcXkBC0HQYJefV+R64JfH/seE56lzlgI5fX+b5pVlxh+HGCuV+ekEzbr62G9/lE+PMG7CAJXDNOQfSUTNuvwSQemPEzN0OO8UofZrondWzD1VyX4orabZHgAz7R8gU0snps7093eWtOmjV5g9r2jXVmIk4SJY/gYdWXa42ZjQ0zfH4CJ0Vpf0GHe5ANrnE3+AWB6fHnqz+hEHKZSBN9eg66buMNoAWGnffEG5Ch0jOnJyNxpbCn2g2+PxTgJmLHKT1Qoko5hJhTryv/X5r9ry6h+353g1HLpi7Cy8FLafpoAwZNCWj/eDo32vjQ+Y/Q+d2A6LJ2IE3P5jk7qxthvrS32QufUtxXDAuD1ZH41fF39U/j56T9ZO5Unwt2ZKIaH2ZHGzIT7fe8fypwiZz4POpKAjWs5orTTReTeSlhATXN7jmTGY+VM4pDkVTBqUdJN05j9SAk3vQtbLhQXwxeLz8N3Kz+EL5ZemKIYzcU7b4KJYfhocrE23fvDt+E/Nv5Vvnv/umUwPKZsTLyYNnZ3WGtZt/PyFrD/+8bmpwFQjMTfWUfavtO3FVVhthCqpeXwfOnL8P2TH8PThS/GXstPFMNjuuOzvz34QznzIukB1cOsn1ZWT/gc0IkRir/nAI2UC6Rdr4SfzKz/cumVuT/jnY05MQyPJkej/2Pnf8K7w9das6y/AxMSbgq0PEyPL/9V9Zvwqvq1BMG4WkQTw/DMeGLlkv/c+LfwqfbevFSPxiYk3AUk6nA8mV8LXy9/G/6y9ouEwNyYBnsnguEx5dn1lXx51qljKCUhYTiAPWYUsHNf/oewVKqE4pgGOSeC4THdSYlkIQMWMSBbKiFhWIBBirNFMfqqafpyoaxU5XHEZDA8VTCXiuGm09MYeU5IGB5QKvjtmPfjTF8T48OnAF3CvQO9MubDlxPD8AkJCVdjMifPJCQkDERi+ISEKUJi+ISEKUJi+ISEKUJi+ISEKUJi+ISEKUJi+ISEKUJi+ISEKUJi+ISEKcItMu3OWCtgaJjGlVYmESIjuhLaGNSlGc0M6u9L06Kz+41i/jp1pj7nLDSkcsb73gduxPCsBMuWPL4o5F0LdGbmxWwozrGmGssjJ4wrYFjIiJ1a2N2HJcHzZAXx8vK+LmjFWBCvY8OQ9kkndM7afNmFXcdUFSasFOdKmoM+Sgri9PRUC68waQvk63YXsJYDPHZm95+dnRvq3PsbMfxWc0Ob7bHBAx1IZ4jx7R83OT1l/biuvPauYU8v6yj97XPXT+2R9BuM/nThmdYKuwuoQnzRCZQBwqDxIZaE+4PIx/qSxSBrx7Ww02TDzpY+g5nZWRFsYaYQ1hafhkppOczP+S623l8nodGuh4PjA60Df3LiNAQio7Ou3Er5SVgsshlmpKSHB4LsBEa0Fyv/ogDZXjoKqicLazda6DK2negXvuDO9s5iLhwnJgRZSJP6D6veN2L4/978D60G2+g0JM1gWLQ0hUHSsTy0Lxppt6Qu9j2dPV9YUOfRUBBC+6Qtxqfj//L0r+GXZ//iD7glqAL3pZE6ZoHwmeeV7f4lI5jE9PeHSD6sR7DRWBd97BrTt06a+l5Ma3SyWKiEH57+FL6sfBUqbP5htAENwDisMvyh9i78sf+rNoGATgD0VSlWtawUq83cVTHcBWJyo23Kx1JqbEMFH6iuVnf44Ie1n7RIxnVB28U2gG7bpkixkBCcbIRxbPT8xeKL8NLazJXm3Zn+2gxPwf59/V+1BQ+6nYUAWLd7vrCowjQ79fD77m+hrk34XLrD7EinP1tDsK43FTps7UuSsxwVjfjT2i/hX178Hzvbzb6bIjYa20G/r73RAhgsVc2KJF8ZoTxbfB4WrIyPqRmmAZjz7L//qfYh/Lb797B7tKPvoQE0M5bcD6s/iyH4LvYHfQdxf6p/CH/b/i9t1cy9+B1FAcH/8uyfw7Jp+HJuh5iHBoxI2dhZln3nWP4cSxeaRvlViyx0+YsJtFfZFb1wNuNwn39m1t9hatoNy2ivtWuCsiUBwvOWrN2+Xf4+fL/6gwnO4TD8tVQfjEnHcMC4SNyvlr+R5HlpFfyy8lIdWrSKU6R4LpXj/GcmmWkIVv38yiT1q8rX4emiL/lLTAAGvQtY4YYNANkOGoJjTTs2oGATSbRHwv0DTU5fY9LKmjNi1mGv+BvWFsyRF77RAoCZ3f91zc+BYimZ4K6U2APwcZeUcmYzPrAXS6ixq9GO0RaKCxfmTL78YN0Z+QcXAMsAIZE/FTcBRYWFtFH/ZDRrAsXoudVuij9cWAwH12J4CovWnrWOWik9CX9Z/Wv408qftTh/tcy+ZBU3242BY9HoVA46lN8qdg5WAczOGt9aGyzzd447x5mAuCk8FnBkEvHweN86YVNSGNcCrcFGFOwtd7t7J9wG2rIpx9AwCuYuWh1aGAQXAO7+RS12TuPQUfZ6TCCosGah4cJcwdj+5Np0RVxJFkK7rgMtzrW0h993XrwFj6HxzxXgPVT5ehreCodkgkHRzPNF98lvA5lqxXmtAvp86WUomzAgaHM7pnSz6MzuCX3kiUIkkhFLQsIwAD1Bb05VV9MV5+L3syFl3C9hr7kjn717uf/hJjs8db/0ej2Gtxd+FRHDlXn2EC/emuEBCwCyoP8zEx74d/hCCJVBiOZQBI3I53jQQEhJTEbK52ZjUdFdTEEWHnQNMhjcX0LDXv1Q5DR3DDongt90zhXnDQsqcVZuXvou1zbxu2kBdb9/OLOr3S95XuwH6BpLc9NM9U/19/L/2bZcQ3m63i3UqJj0fg1+j33u9Hazel9Tw1PBU5nmwxoawXwj6Ldo97SiZ9/2gueqcXhljchnovy8q8IzZ6FUKIdKuWrC6ImskEUzvXh/tvA8rM0/vby81l48v59J1KB2YGZx8Lzczz3gWoaXYrSVz9dB7CzV76KbXwC1Da/M/Itt42UlLnKx78f3+WOcka/DTesTz7/pdZdB97IX/nrDNDv++PvDN6bdX4eafeZ7f1YsM/9fTJ/nv2RlxD3g6BgP0Nf26Ua0cy2GR2NWSyvackcMf0kBrwtMGIbliL6uzj/TZydYH98kUokpRPSSgBzfY/qzWT8BOdagZ0gE5qe+T8qr4ae1fwo/rP0c/rz6F8UIsCJI6oidyXv3GUf+DPP5Cbyw2yzBF87nN77n/h9rH3RIOlt5aOTYqTAW2wpz/VZzy6T4RzPdNlRehgj7GV/XUYZYxxOvY9Pq1Wg3PDpr30v4xDJnLxibchPwYey3ofLthL2jHT0LzcFIBUFLyrtt7cT9Jaiy+3A9dcQ947d48B3PvUjwjhrUHlYv6o0/TN05mh1rQ/ssP5h+t1cetKmSfKwPaQO1g7W5H/zt+QOxzW4Mexx9x/22TKv/sfdr+MfO3+zvTZWJZ6MQYrmhAQ/KXdzu1CDWl2sYHSAfZr3+SQE+lmZvEtwz+og0cxnm/q8h+/tCoCExmzkw5QdpTCr54fCdGp3CcU4ckiEyz3Y9/YDJMe+5rzrQGhwTqNY2BjSm2TUChsHpIM6hokRHeWfjCRqRkszZbxJCpu2jpkPywYwUdXXBtLy94jMQHDQcwyEMhRDw4zfKS+MeGPMT3YfJt4iY2u/147oYhnNwFzifIUiYjHOQ5AgI7nt0ylhykJtBuSPUcdYx9U5NQoE6wrhI/gMTGjwDAUbgi3aW22TlF6NaZyOEaB8N41i78Gzam2dQD8qL+UgZiPDG9qXMtNXh0aEizDybeyFkOJd7tM/suWduVuq5VoFLLaM+UDf6iTbfaHxSGwPuRT4E7hX9gNvVD+qH4KP89A9MQMVpP4LCjPCgdCID00bQ257VhWg2bcL1vCOUj1ECBtqFOvDKMw3n7TM0bHWnnLQF13Mt5/DsOO6dbwPKxRg8jMY9uCe/k0GIVblmLupCcUFtAD1smN8OHVFWSk4/8Dv18ntbu9kPCAIUDsPVCHQYl3bDoibhaNncaO4Jc9NGojXrf/qRfqUc7soy3Hm5Dr8Ww+dBQQfhNgwPOI8CQ9AQLY0Jwby3e707eKOKdUwS8lxyANgZluGQXTuXiCdSkntDNNvWGGJAk6iM6/K3lSJ8Xf2GB0my8h0Nti5NyNDdu7BrDc096DgakOG9N/t/6OD5CB46g0QLOqFqxEs96VDK+MbKxfPUEaZxYWKIn87ivpFo0B50HAEcNDHtBeHxDOpN2dAydJ7SUOecYGlbOhdm5ryPtfcyE3daW7ofxMG9+I4yEBhC4yH8KC/3Q0CuNz6qvIxcUB+et2F/c28EBNlwDK06owwW7BeBPrxvhgcwB/dG4FGXTXsWf9NH9Dv1p5+hPYLDJvJUBpiIvoVuqP+H+ju1J9dBT9stu7axpbLMzlqZZ4mck0sKY3o7XMnwC8+sv8tq+x3rA4Qr+Sn0A+dxvxj/itpYf9uLPsgzPM+FdqpmVfOOpUu/vzn4XbRDLgD8QruBRXNh6WeU32W4McNfhNsyPKDCdMamNeQfe7+dNxQEhAhEozIy4FKsoEY5MT9mzhqLBibiz/P/sfc3SUEYk4woop64IWRp0Yk8AwaD2elsGgttwHDigjEGxAbhW+H1DA4IgPKplBkDUQaEC+dDCmiNvCkoQrAXuQncF23BfXg2AguTm7qtGQOQ0wChwAhyL4wx6XykdcEIifoR/Nk3QfJm/3fdg0wsN/3PVJal0pLaCKagHJSA5/NsfsPeIQMOQoUhl8psm/RCIy7UgXaQtWFtDrHRV9wLYrwuqPMghqcc9Bv1g3i9jJ2eAwbFLUIbYh0piq3rugwvl8++Z7waIQ2z0+er80+tf/9k2vWZ2gIFcGAWGc/n3iRdMeRFX9LnUcjjAzPE9u3K99ZGVZUd6xHa4b5cWyqUVIbrMjxJXriotBvDdtAgfQnD8x3tSqCa5CPOpW4IAOhmkIbHOuN32gK65DcgoWgvCQ34zGilZOcRtyKedRmu36P3DCQTjQ5BQ/hK07VGp1EhGDTEi6WX1vlfqbHIQmL8UpW0Bqcz0FgyUclvtmtp6DxoRIhCBGbn6Rl2Hc+hkzmfZ0EIpHEiSKLEpHH9OkzxXX1eLCyJWVcX1hQviOfpfvjGJgioE3VwYbPnFoVpJL7jGWSSUa8XRgQIRxgWhiEeQH1IygCUDSLD7I8ECZtBEKStxuQmCIq2grhIZuE8iA5Nzv2OT1vW6bPWbqZxjVlgRAgONwFC5tloExccdwf15P60nbtQZi31H6YNETrqM82F+Bzch/ZEG2/UKOOW2rZSqqjetCF0Qd0ROmhANCH3h/nxc2H4dbsWpQDz0HZYHcR/sEJo15hUs24WGzSCALk2jPmcsctSDNAnAs9/sr/sNwQIZa6WqxJGCOooUPKgvlgz4gFozfrpiQkTZbdmw+KcQz2gd8rqlpFbDhdhJBg+NhISFU2ctZHe6RQYj0Ackvyb6p/Cd2zOb4QNw3yz/J1MZxiVxh3UeCBqGrQqz4qgMWEmiAlG/2H1p/DNynfhu9UftD0wJq7OUwfAdOYrmi/8svK1lYNzvs3KsKrzAPfEGkALQ8T47VgsMDrmPN+jfbEAIDQYkzogPDDN8FEVlDEC5Vq0YmmWeIBLez3DXlhQTEb51spL23ydlQXCJ0GKxCikP7EGmADGs6KFjr0DhB/ty328DY5kHmM6ttpkj11OPNeCFRfhAUO7C/T6swN3CAY+VuxjMMND+LhwYmITuLQrmjUKtrK1D4z7hdW9YPWhPw/sPAQYZjvxCp6P0KFWWEcI086JmfD2wloEjLbA+HumcTmXfrspvNms4kaLg1qQvrsK0Ct8gUB+ajT7rfUr1gjZqswvcM1PuzBq5Rl8tNFVGBkNrzaIRw58FCOfN8CCtNmPaz+H71d/lMSDEWjEi5i9C7+POiMHCB9iWTTzF40OAdCg+IAxCBKfj6Rl+A9C41yuZbTBmdHhwgHp69NFIfjNzOfKdzb3JLgoK8OYEAJGKHFOxyQ3GYRumrqVkK+faww/jFzFtDIZF54rk/HHpz9LKGIBURfui1kMiIm4S5C5IFmRYHCyHnkmzHcdAroK1JFnL5tgwn1BWPcfCCnMXE+v7SVJaky5aAM0GRobhoz19t/MMuBlDAwQclyIAKPNEXb0E8yjfrNncF+0t0Zezrx9BfsBpj8yCw1BgatxU6ifuF23q4VYF//rCtg9oAfKjNXC3/Qx7if9SO0jGJij7hqgi515AUaH4QdBZffGgXAAjAURYwqj4TGd1IE08i3AfWHceWtYJtxEgtPogR35+/InxAuD0/gqk305Z2WCoCK8rFknoEWMWDGrMcn5TLUgUsx+huSIe0DMmP1oOI+2kmpp59j3SPE8umXyd3xUXpSBRCMsB9oHTYCAor2wQAg20l6RsWR9iEmcSCB6hA+MQJnx/e8Kta+1I89FOONy9B9YXVgq9G2ekEFkRAQiQgjmpWwAKwqmpg1xA2lH3DQ91eqH0OV3BBvmNc/BjaE9+Iz1FjW92iFrVp6JlSVXwJ53G9BFfvhNIyNmjzB0/xoEfqWPXFgTR/B2oT1RSP30Tj14hj/lYow2w18AKosUh+mGhlwDxsbsb9SL4Gd93tRcTifQGS0CicZIwAm1JQ3CQXSY6DK+c8N8dMz84owJH3spBiCmvBoiBiMMDcdJE5hFlEWPv17+5tzHR5tC8JGx0Q29yLRmZgE8NmhZgrT497SFt6mb3mhg3BAfuXGNzDkIawQbAhqtjt+MO/iNuT7PMmWxQuJXcUnnu//r8Pvbd2bu85yr2egCDLrMiMJufwNwcu8FF5Pl1fQ6lgwPIO7rMuSjAqa3F0wPEQE+R9+L72HOglkXaCB8cVwVpkQSoIQg8+7CdUC78AIwP/7/svn0aHysAIbofDrmjpg+livCSexGVHmPoB5WFntz28jrpfbUNz6rDg2IhYa7ha+r5Ctrw1fVb63uK2oHhByR8Sd2Drch0MUQHSMUURgL9ps/1V7WNrEtbwJdMeAyhMkwyPa2tD+2DD9uEBPmOomOh+kVK5gzP83MTPzvl9WvpIVeVb9R0I0gHn7bbSGXwrQX0dwjI2oI+xitjuZyzvYjA+XSob/9u5FArpxZ6VQ3rDyGonBdcFnijEzaEOHpox+VzCXw2WgcDK/JuoGh1S+fM5AYPvv7utC97KLzts3DftLPWRs/BhLD3zOclMy/xsxGE+kbaAH//kSMThRWow6ZT4u5CeHK5yZ5xK69DSBmdx22tSjFm4PfpNnxYV+Yef9q+evzgFAPvIiGxyHKXniwkiJpBCcDJePbopntmOv45rENdZgb86S8pt+wahi7ZliOdNe/7/xNQ6u4O382a4p4B20SK35XhjwX7hw94HP/dw+LxPD3DIgSk53kjgX8RVjdfFECQoz3kiRC4k70FU9PCSIxjOf53XGdwJuCe2FB7DU90Yi8e+6HX88qRB7s9OBgPxHy6bE00GewYlAUrBz8cnxyyu3CjGHELQ3VKeHr1MesOTpWd3x8hlEZliQaz/wLBB7ncY8YEEMgfNYK9oEWuG0r6Lq+i+8qSIaBITP8gMpYw0WtNsq4bgm9w/rOvvRi00JGWB4l90wzzof5ID7G25lwQ7CJgB2ES8QZv7JmhNpt0osfMoiI+I5sL+6hYGCnLkLnLvA4vxOJhnF6g4Ld53COXgPufxGk2LK/z5E976r76LrPLvavYHZ88JhUBIiBMOxG3gBpucqws/ZjzJ1xf0ZFEJyK5Ns7Qo/vIqgbZr2b9lmAMldElTdX7p7i95ez7/NADa8bfN4OfKQsg9H/oLthKAxPYWNxXWM44p8XV+ZzIBxuIyCue50a97MG906IV3c7+PPzurj6WZwxZ2Yo2VQkhDxd9EQbygmj4Uu+3X8dft393/Du4K1MTgjXs922JADyZmz/M2O59fWA4vAMtBzWBH+j8fHhyfNmll3DhIHSNCOxG+gz7gkD8D3XDbr3IIgO1EbdC3rLOPhGfo2u5oO+0+n6w6/xIUfPtsTVwTliKC0m9GCq05akZ5MXTzsSlJy1y7FqqLs/Qc6U4hkSrK09WQEIBS9HBhXJzuaw8znOqxXfB8GuEx3m3oWea7JyWPnVvtlze54vdB/Y89tFz7/o+xyGwvCYnUhbCCRGNXlRSCeyy5M4qDy/q/L+hcBCf3zgvjFFdRDoeNY153ma8cXLCMWlrEnxM7+3ynJKrnK3jC6g7BzTducEzmd7Hs/kkZwH4fHiWXHISt2mZ3Nt954RZLRxLueQ3edR+G+V349pyrNpGwiTjDMRrb2TFspabkSUeTr10vN4BuWlGXJNETV1HpSNsmBdMDzFWDh1wKdHoPA8JihhUQA3a+OzXFDwGwKhh9gGIF5D+8E4vMe2oLxE1GnPnj42cB0vGJB2Ultl5eY6/j6x62gj+0rZiV8sPQ/KJFTyU5mb6Pcda0PmKbDyLXMOGKJjrYSFwpKu85GQgr0XuETlZJYjsQ3mgJCyLDfHfkdA8DyeDw1gjamPKSNJOnaDrHZZPbJRA6sbB9dp2Dg7dD8D9UTg8mzalXeCqU53JnStvDxY91WTd9s1tpOfa5+yMqig0ITOJXOilw76MRSGpxIk90OQNCrEDKHxN0CSQmgXgUoRNYUhuQaNiG/FMAufqaw0kf3eDyrO8xlH9Wt9tRuujWmxEC8CCaJD23E/ysiBic2QDR3L4VlvdIzPldawFueqPkXVkbogRGh8mBpzkbLFe1J3rlNEXGXzaHKcyMPSXowBEzTiewidISLMeBJwSKSB2fG16U/KxTl0cGyfbvt6jnlklgg0IITGsN6ThVVNookJHBAwZi/Cz8eoqwpa8Tf35jqYiMkf3DfPpIPgTO2pxCQTcX2+ffmddlI9rN0yalbdJDCtHJQFaMLIeRuaZrbfY5495ScFmdGL55WXnh1p9aMNlShjbgsamzaEBfgNZtd4vNEDi6TwHeWijekb5kZAFwgPAn0E+bDCSiRXifa8v0Uf9gzo1PvB6Yx3mBqmp4zci9+VJWfPpl0pNzkVgHZlAhB1gqalWCSwfcIM9RYN2fUMNyJkohDVYZ9pP8Ug7DxvY88sFL11aEcXCIMwlNly5DZjgkL4FIR16iAeL7gTEN/R2NJQOVAwKkLUFB+MytH4nhFm9yiU9DeNhg9HA0dwLS+uI5ONBuS5nM+5dLR3njWKNSBExbgrQTIvp51T9N8pF8M3pDJSQiXEmAaAIHk2Qz/cCyajI6rzy1YvBMWx8uMRJHzWPe1cdbLq7Su5Qjx0JkQAo3MviID7cQ0pw6zO82LpVfhm+XsJBJ5JHSUMMkKGELhW9YuH6mkCMqtHBM9nii1lIYjF8+OiIyTfvKp8dT7WHzUbjMG9CTAyQSlG8fv7LSIKGQiXPsCUxuKjLLFsTpQ+q8/7w+eDW+UkKGAAroV+VEbVy9uwaOW3h6vcSn3Gl7e2IqfA72uCwQh+3s5HmyMQXlZfaa7D88UX3ibmFlE32pJ6kpnJvZg4w7AdQ6DPK19KIFMbtYNdR7u4gPDAHtqdmAH0Q/nz5aSecUIXB7/TNgg4MZ+0N+3lCg5XjXaFDpkjAP0gVMq0m9WVtvN7Ge0bD/B8BDU5FLh6IJ7DO/fib8pdnltQfw3qs1vsLfc5KDAFgSD7tYGI3iqAthq0vhyNgnakwkRS0Tx5cDrSFqaB+Kh4hNjdiu+BrkP5YnnwLAiImWEwFGVjQoSsCRMyEZwHs6JVeQ4SEyKsHR2oE/JNJJ981jSNmZQQG/dk9hVaoN8CoRMgwjgNkrbgfNeEvhBHmxV7McOsDBANAoFZeAgoEam9OA8riTpmfCLE+jlhopm6UzmBCA4NYYxEfdAkoCs0IOSStA7tTjsi4KIw5L4yc+1zf7/l4X3oKxLBuP1WAddyj4q1A9NIuSdQ/1ndFWSz+nF9vq15JAxOWaAfBCXgKvqlZdf5rERfXQh3DgZHaLtQ9faI7YBWxWqjvrQt/Uc70OacxzkE9VAc3I92gplpB85FqEVrNbYl4FoJNPUDQm1B9MVzoDcFZK1t+Q6BSvkQLpQxMjLWCfc9O/X6c0/KSFtpDoCVU21s59FeuKAR8flOB/Z85vL3xH66GArD0zh03IWwJ2BKIvUoXD8ogvsw2RcD4UTTD5GMiOuSauSef/lz/AcxmpXpwjrxKKsHL//IFxc938/jnoMAg+lKu1VccaYfal/qeFG57QbcP+ZbXwSvU9dHziP+BviN103B9WrfC+DP9X7I41L6yRonlunz9vFycw8/g3YYXHbO44zL4OxgZ9q7nql7eWxDz7HXxXfw82M/6nxTAvLVTcij7WFMWT/GwJEmdO8L+pfyUiJ3CbwMvAaDsy+mNTAUhqcAPOgyXOec+0RspGGX4e51p1xOHHcp27XKkXX1RQwRwXlXnXNb3Lael5Up/nbXNoy46FnXu7/3Zx4wMz56/FrC2V7X7Q8wrLoNheETEhLGA5fbgAkJCROFxPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVOExPAJCVODEP4//kJgV52WzjEAAAAASUVORK5CYII=" alt="Trimera Health logo">
                <span>Trimera Health</span>
            </div>
            <div class="trimera-topbar-app">📋 Trimera AI · Clinical Intelligence</div>
        </div>
        """,
        unsafe_allow_html=True,
    )



def render_sidebar_brand() -> None:
    st.sidebar.markdown(
        f"""
        <div class="trimera-sidebar-brand">
            <img src="data:image/png;base64,{logo_b64}" alt="Trimera Health logo">
            <div class="trimera-sidebar-brand-text">
                <div class="trimera-sidebar-brand-name">Trimera Health</div>
                <div class="trimera-sidebar-brand-subtitle">Trimera AI</div>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_trimera_header() -> None:
    st.markdown(
        """
        <div class="trimera-page-card">
            <div class="trimera-kicker">Trimera AI · Clinical Intelligence</div>
            <div class="trimera-title">📋 Documentation QA</div>
            <p class="trimera-subtitle">
                Grounded review using payer, AMA, CMS, and Trimera authority documents.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


APP_TITLE = "Trimera Documentation QA"
MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")

ROOT = Path(__file__).resolve().parents[1]
REFERENCE_DIR = ROOT / "reference"
MANUAL_PATH = ROOT / "Trimera_Documentation_Coding_Standards_Manual.docx"

REFERENCE_FILES = {
    "AMA": ["2023-e-m-descriptors-guidelines (1).pdf"],
    "CMS": ["mln006764_evaluation_management_services.pdf", "ai.rules.emcodes.docx"],
    "BCBS": ["cpcp051-12-22-2025.pdf"],
    "UNITED_COMMUNITY": ["UHCCP-Evaluation-Management-Policy-(R5007).pdf"],
    "DOWNCODING_RISK": ["OH.PP.066.pdf"],
    "TMS_CMS": ["Article - Billing and Coding_ Transcranial Magnetic Stimulation (A57528).pdf"],
    "TMS_UNITED": ["transcranial-magnetic-stimulation. united.pdf"],
    "SPRAVATO_MANUFACTURER": ["SPRAVATO_Access_Coding_and_Reimbursement_Digital_Guide.pdf"],
    "SPRAVATO_UNITED": ["PA-Med-Nec-Spravato.united.pdf"],
    "TRD_BCBS": ["major-depressive-disorder-doc-code-guideline.pdf"],
    "TRD_INTERNAL": ["ai.rules.trd.docx"],
}

PAYER_AUTHORITY_ORDER = {
    "Medicare": ["CMS", "AMA"],
    "BCBS": ["BCBS", "AMA", "CMS"],
    "UnitedHealthcare Community Plan / Medicaid": ["UNITED_COMMUNITY", "AMA", "CMS"],
    "UnitedHealthcare / Optum Commercial": ["AMA", "CMS"],
    "Aetna": ["AMA", "CMS"],
    "Cigna": ["AMA", "CMS"],
    "Humana": ["AMA", "CMS"],
    "Other": ["AMA", "CMS"],
    "Not specified": ["AMA", "CMS"],
}

MDM_RANK = {"none": 0, "minimal": 0, "straightforward": 0, "low": 1, "moderate": 2, "high": 3, "unclear": -1}
E_M_RULES = {
    "99203": {"mdm": "low", "time": 30, "patient": "new"},
    "99204": {"mdm": "moderate", "time": 45, "patient": "new"},
    "99205": {"mdm": "high", "time": 60, "patient": "new"},
    "99213": {"mdm": "low", "time": 20, "patient": "established"},
    "99214": {"mdm": "moderate", "time": 30, "patient": "established"},
    "99215": {"mdm": "high", "time": 40, "patient": "established"},
}
PSYCHOTHERAPY_RULES = {"90833": 16, "90836": 38, "90838": 53}
TMS_CODES = {"90867", "90868", "90869"}
SPRAVATO_CODES = {"G2082", "G2083"}

FACT_EXTRACTION_PROMPT = """
You are a documentation fact extractor. Do not decide whether any code is supported.
Extract only facts explicitly documented in the completed note.

Return valid JSON only, with no markdown fences, using this exact schema:
{
  "patient_status": "new|established|unclear",
  "total_em_time_minutes": null,
  "time_statement_text": "",
  "em_time_separate_from_other_services": "yes|no|unclear",
  "problems_level": "straightforward|low|moderate|high|unclear",
  "data_level": "minimal|low|moderate|high|unclear",
  "risk_level": "minimal|low|moderate|high|unclear",
  "overall_mdm_level": "straightforward|low|moderate|high|unclear",
  "problems_evidence": [],
  "data_evidence": [],
  "risk_evidence": [],
  "prescription_drug_management": "yes|no|unclear",
  "psychotherapy_separately_identifiable": "yes|no|unclear",
  "psychotherapy_minutes": null,
  "psychotherapy_intervention_documented": "yes|no|unclear",
  "psychotherapy_response_or_progress_documented": "yes|no|unclear",
  "psychotherapy_evidence": [],
  "longitudinal_relationship_documented": "yes|no|unclear",
  "longitudinal_evidence": [],
  "base_em_code_or_level_explicit": "yes|no|unclear",
  "prolonged_time_separately_attributable_to_em": "yes|no|unclear",
  "diagnosis_documented": "yes|no|unclear",
  "medical_necessity_documented": "yes|no|unclear",
  "baseline_symptom_scale_documented": "yes|no|unclear",
  "informed_consent_documented": "yes|no|unclear",
  "tms_session_number_documented": "yes|no|unclear",
  "tms_motor_threshold_documented": "yes|no|unclear",
  "tms_motor_threshold_method_documented": "yes|no|unclear",
  "tms_coil_placement_documented": "yes|no|unclear",
  "tms_intensity_documented": "yes|no|unclear",
  "tms_pulses_or_protocol_documented": "yes|no|unclear",
  "tms_treatment_delivered_documented": "yes|no|unclear",
  "tms_tolerance_documented": "yes|no|unclear",
  "tms_adverse_effects_documented": "yes|no|unclear",
  "tms_safety_assessment_documented": "yes|no|unclear",
  "tms_repeat_mt_reason_documented": "yes|no|unclear",
  "tms_prior_and_new_mt_compared": "yes|no|unclear",
  "tms_plan_documented": "yes|no|unclear",
  "spravato_indication_documented": "yes|no|unclear",
  "spravato_rems_documented": "yes|no|unclear",
  "spravato_dose_mg": null,
  "spravato_self_administered_under_supervision": "yes|no|unclear",
  "spravato_pre_dose_blood_pressure": "yes|no|unclear",
  "spravato_observation_minutes": null,
  "spravato_periodic_vitals_documented": "yes|no|unclear",
  "spravato_sedation_assessment_documented": "yes|no|unclear",
  "spravato_dissociation_assessment_documented": "yes|no|unclear",
  "spravato_respiratory_monitoring_documented": "yes|no|unclear",
  "spravato_discharge_status_documented": "yes|no|unclear",
  "spravato_transportation_documented": "yes|no|unclear",
  "spravato_followup_plan_documented": "yes|no|unclear",
  "important_ambiguities": []
}

Rules:
- Use null for an absent number.
- Never infer time from appointment start/end unless the note explicitly attributes that time to the physician/QHP's billable E/M or psychotherapy work.
- Do not count time for another separately reported service as E/M time.
- For MDM, apply the 2-of-3 framework, but extract the level rather than deciding the billed code.
- Prescription drug management may support moderate risk when explicitly documented.
- For TMS, distinguish initial mapping (90867), routine subsequent treatment (90868), and repeat motor-threshold determination (90869).
- For Spravato, extract the administered dose and the documented monitoring/discharge elements. Do not assume a 2-hour observation from appointment timestamps alone.
- Keep evidence short and quote or closely paraphrase the note.
""".strip()

EXPLANATION_PROMPT = """
You are Trimera Documentation QA's explanation layer.
The Python rules engine has already determined the code-level outcomes. You may not change those outcomes.

Use the completed note, extracted facts, governing excerpts, and fixed findings to provide concise, useful feedback.
Rules:
- Never override SUPPORTED, BORDERLINE, or NOT SUPPORTED.
- Never invent facts.
- Explain material gaps only.
- Distinguish true billing deficiencies from optional documentation-quality improvements.
- Cite authority using the supplied source label, such as [AMA], [CMS], [BCBS], or [UNITED_COMMUNITY].
- Do not recommend a different code.
- Keep each code explanation to two concise sentences.
- Give no more than three documentation-quality improvements.

Return valid JSON only:
{
  "code_explanations": {"CODE": {"support": "", "deficiencies": []}},
  "quality_improvements": [],
  "final_assessment": ""
}
""".strip()

FOLLOWUP_PROMPT = """
You are Ask Trimera inside the Documentation QA module.
The code-level result was produced by a fixed Python rules engine. Do not change or contradict the fixed result.
Answer using the completed note, intended billing, payer, extracted facts, fixed findings, governing excerpts, and prior follow-up messages.
Never invent facts. Clearly distinguish documented facts from missing items. You may explain a result, identify the exact gap, or draft concise provider education or an internal billing note. Do not guarantee payment or audit success. Keep the answer practical and concise.
""".strip()


@st.cache_data(show_spinner=False)
def read_pdf(path_str: str) -> str:
    """Read a reference PDF only when a QA run actually needs it."""
    path = Path(path_str)

    # PyPDF is materially faster for the policy files in this app.
    try:
        reader = PdfReader(path)
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text
    except Exception:
        pass

    # Fallback for PDFs whose layout PyPDF cannot read well.
    with pdfplumber.open(path) as pdf:
        return "\n\n".join((page.extract_text() or "") for page in pdf.pages)


@st.cache_data(show_spinner=False)
def read_docx(path_str: str) -> str:
    doc = Document(path_str)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return "\n".join(parts)


def read_document(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return read_pdf(str(path))
    if path.suffix.lower() == ".docx":
        return read_docx(str(path))
    return path.read_text(encoding="utf-8", errors="ignore")


def split_text(text: str, source_label: str, chunk_size: int = 5000) -> list[dict]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    chunks, start, index = [], 0, 1
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            boundary = text.rfind("\n\n", start, end)
            if boundary > start + 1000:
                end = boundary
        body = text[start:end].strip()
        if body:
            chunks.append({"source": source_label, "chunk": index, "text": body})
            index += 1
        start = max(end, start + 1)
    return chunks


def authority_order(payer: str, codes: list[dict]) -> list[str]:
    """Return only the source groups needed for this specific review."""
    order = list(PAYER_AUTHORITY_ORDER.get(payer, ["AMA", "CMS"]))
    code_set = {entry["code"] for entry in codes}

    if code_set & TMS_CODES:
        if payer in {
            "UnitedHealthcare Community Plan / Medicaid",
            "UnitedHealthcare / Optum Commercial",
        }:
            order.append("TMS_UNITED")
        order.extend(["TMS_CMS", "TRD_INTERNAL"])

    if code_set & SPRAVATO_CODES:
        if payer in {
            "UnitedHealthcare Community Plan / Medicaid",
            "UnitedHealthcare / Optum Commercial",
        }:
            order.append("SPRAVATO_UNITED")
        if payer == "BCBS":
            order.append("TRD_BCBS")
        order.extend(["SPRAVATO_MANUFACTURER", "TRD_INTERNAL"])

    order.extend(["TRIMERA", "DOWNCODING_RISK"])

    # Preserve priority while removing duplicates.
    return list(dict.fromkeys(order))


@st.cache_data(show_spinner=False)
def load_reference_library(
    payer: str,
    code_signature: tuple[str, ...],
) -> dict[str, list[dict]]:
    """Load only references relevant to the selected payer and codes."""
    codes = [{"code": code} for code in code_signature]
    categories = authority_order(payer, codes)
    library: dict[str, list[dict]] = {}

    for category in categories:
        if category == "TRIMERA":
            library[category] = (
                split_text(read_document(MANUAL_PATH), "TRIMERA")
                if MANUAL_PATH.exists()
                else []
            )
            continue

        category_chunks = []
        for filename in REFERENCE_FILES.get(category, []):
            path = REFERENCE_DIR / filename
            if path.exists():
                category_chunks.extend(
                    split_text(read_document(path), category)
                )
        library[category] = category_chunks

    return library


def parse_codes(raw: str) -> list[dict]:
    entries = [item.strip() for item in re.split(r"[\n,;]+", raw) if item.strip()]
    parsed = []
    for item in entries:
        match = re.search(r"\b([A-Z]?\d{4,5})\b(?:\s*(?:x|×)\s*(\d+))?", item, re.I)
        if match:
            parsed.append({"raw": item, "code": match.group(1).upper(), "units": int(match.group(2) or 1)})
        else:
            parsed.append({"raw": item, "code": item.upper(), "units": 1})
    return parsed


def query_terms(codes: list[dict], payer: str) -> list[str]:
    terms = set()
    for entry in codes:
        code = entry["code"]
        terms.add(code)
        if code in E_M_RULES:
            terms.update(["medical decision making", "two of three", "problems addressed", "data reviewed", "risk", "total time", "prescription drug management"])
        elif code in PSYCHOTHERAPY_RULES:
            terms.update(["psychotherapy", "separately identifiable", "time", "90833", "90836", "90838"])
        elif code == "G2211":
            terms.update(["G2211", "longitudinal", "continuing focal point"])
        elif code in {"99417", "G2212"}:
            terms.update([code, "prolonged", "total time", "99205", "99215"])
        elif code in TMS_CODES:
            terms.update([code, "transcranial magnetic stimulation", "motor threshold", "coil placement", "treatment parameters", "tolerance", "adverse effects"])
        elif code in SPRAVATO_CODES:
            terms.update([code, "Spravato", "esketamine", "REMS", "56 mg", "84 mg", "two hours", "observation", "blood pressure", "sedation", "dissociation", "discharge"])
    terms.add(payer)
    return sorted(term for term in terms if term)


def score_chunk(chunk: dict, terms: list[str]) -> float:
    text = chunk["text"].lower()
    score = 0.0
    for term in terms:
        term_lower = term.lower()
        if term_lower in text:
            score += 4 + text.count(term_lower)
        elif fuzz.partial_ratio(term_lower, text[:12000]) >= 90:
            score += 1.5
    return score


def governing_excerpts(library: dict[str, list[dict]], payer: str, codes: list[dict], limit_per_category: int = 5) -> tuple[str, list[str]]:
    order = authority_order(payer, codes)
    terms = query_terms(codes, payer)
    sections, used_sources = [], []
    for category in order:
        chunks = library.get(category, [])
        ranked = sorted(((score_chunk(chunk, terms), chunk) for chunk in chunks), key=lambda item: item[0], reverse=True)
        selected = [chunk for score, chunk in ranked if score > 0][:limit_per_category]
        if not selected and chunks and category in {"AMA", "CMS"}:
            selected = chunks[:2]
        if selected:
            used_sources.append(category)
            sections.extend(f"[{category} | chunk {chunk['chunk']}]\n{chunk['text']}" for chunk in selected)
    return "\n\n---\n\n".join(sections), used_sources


def extract_pdf(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    return "\n\n".join(f"[Page {page_number}]\n{page.extract_text() or ''}" for page_number, page in enumerate(reader.pages, start=1))


def clean_json(text: str) -> dict[str, Any]:
    cleaned = re.sub(r"^```(?:json)?", "", text.strip(), flags=re.I).strip()
    cleaned = re.sub(r"```$", "", cleaned).strip()
    return json.loads(cleaned)


def yes(value: Any) -> bool:
    return str(value).strip().lower() == "yes"


def unclear(value: Any) -> bool:
    return str(value).strip().lower() == "unclear"


def calculate_mdm(facts: dict) -> str:
    ranks = sorted([MDM_RANK.get(str(facts.get(key, "unclear")).lower(), -1) for key in ("problems_level", "data_level", "risk_level")], reverse=True)
    if len(ranks) < 2 or ranks[1] < 0:
        return "unclear"
    return {0: "straightforward", 1: "low", 2: "moderate", 3: "high"}[ranks[1]]


def evaluate_em(code: str, facts: dict) -> dict:
    rule = E_M_RULES[code]
    calculated_mdm = calculate_mdm(facts)
    mdm_supported = MDM_RANK.get(calculated_mdm, -1) >= MDM_RANK[rule["mdm"]]
    time_value = facts.get("total_em_time_minutes")
    time_supported = isinstance(time_value, (int, float)) and time_value >= rule["time"] and yes(facts.get("em_time_separate_from_other_services"))
    patient_status = str(facts.get("patient_status", "unclear")).lower()
    patient_match = patient_status == "unclear" or patient_status == rule["patient"]
    if (mdm_supported or time_supported) and patient_match:
        status = "SUPPORTED"
    elif MDM_RANK.get(calculated_mdm, -1) == MDM_RANK[rule["mdm"]] - 1 or (isinstance(time_value, (int, float)) and time_value >= rule["time"] and not yes(facts.get("em_time_separate_from_other_services"))):
        status = "BORDERLINE"
    else:
        status = "NOT SUPPORTED"
    deficiencies = []
    if not mdm_supported and not time_supported:
        deficiencies.append(f"Neither {rule['mdm']} MDM nor clearly attributable {rule['time']}-minute E/M time was established.")
    if not patient_match:
        deficiencies.append(f"The note identifies the patient as {patient_status}, which does not match this {rule['patient']}-patient code.")
    reasons = []
    if mdm_supported:
        reasons.append(f"Calculated MDM is {calculated_mdm}, meeting the {rule['mdm']} requirement.")
    if time_supported:
        reasons.append(f"Documented separate E/M time is {time_value} minutes, meeting the {rule['time']}-minute threshold.")
    return {"status": status, "reasons": reasons, "deficiencies": deficiencies, "calculated_mdm": calculated_mdm}


def evaluate_psychotherapy(code: str, facts: dict) -> dict:
    minimum = PSYCHOTHERAPY_RULES[code]
    minutes = facts.get("psychotherapy_minutes")
    separate = yes(facts.get("psychotherapy_separately_identifiable"))
    intervention = yes(facts.get("psychotherapy_intervention_documented"))
    time_met = isinstance(minutes, (int, float)) and minutes >= minimum
    missing = []
    if not separate:
        missing.append("Psychotherapy was not clearly separately identifiable from E/M.")
    if not intervention:
        missing.append("A specific psychotherapy intervention was not clearly documented.")
    if not time_met:
        missing.append(f"At least {minimum} psychotherapy minutes were not documented.")
    status = "SUPPORTED" if separate and intervention and time_met else "BORDERLINE" if sum([separate, intervention, time_met]) == 2 else "NOT SUPPORTED"
    return {"status": status, "reasons": [f"Separate psychotherapy: {'yes' if separate else 'no/unclear'}; intervention: {'yes' if intervention else 'no/unclear'}; time: {minutes if minutes is not None else 'not documented'}."], "deficiencies": missing}


def evaluate_g2211(facts: dict, payer: str, codes: list[dict]) -> dict:
    base_present = any(entry["code"] in E_M_RULES for entry in codes)
    longitudinal = yes(facts.get("longitudinal_relationship_documented"))
    if payer != "Medicare":
        return {"status": "BORDERLINE", "reasons": ["G2211 is Medicare-specific in the supplied CMS authority."], "deficiencies": ["The selected payer is not Medicare; payer-specific coverage would need confirmation."]}
    if base_present and longitudinal:
        status, deficiencies = "SUPPORTED", []
    elif base_present and unclear(facts.get("longitudinal_relationship_documented")):
        status, deficiencies = "BORDERLINE", ["The ongoing longitudinal relationship was not clearly established."]
    else:
        status, deficiencies = "NOT SUPPORTED", []
        if not base_present:
            deficiencies.append("A qualifying office/outpatient E/M base code is absent.")
        if not longitudinal:
            deficiencies.append("The continuing focal-point or ongoing-care relationship is absent.")
    return {"status": status, "reasons": [f"Qualifying base code present: {'yes' if base_present else 'no'}; longitudinal relationship: {'yes' if longitudinal else 'no/unclear'}."], "deficiencies": deficiencies}


def evaluate_prolonged(code: str, units: int, facts: dict, payer: str, codes: list[dict]) -> dict:
    base_codes = [entry["code"] for entry in codes if entry["code"] in {"99205", "99215"}]
    base = base_codes[0] if base_codes else None
    total_time = facts.get("total_em_time_minutes")
    separate = yes(facts.get("prolonged_time_separately_attributable_to_em"))
    if not base:
        return {"status": "NOT SUPPORTED", "reasons": [], "deficiencies": [f"{code} requires a qualifying 99205 or 99215 base service."]}
    if code == "99417" and payer == "Medicare":
        return {"status": "NOT SUPPORTED", "reasons": [], "deficiencies": ["For Medicare office/outpatient prolonged time, the supplied CMS authority directs use of G2212 rather than 99417."]}
    if code == "G2212" and payer != "Medicare":
        return {"status": "BORDERLINE", "reasons": [], "deficiencies": ["G2212 is Medicare-specific; confirm the selected payer's rule."]}
    first_threshold = (75 if base == "99205" else 55) if code == "99417" else (89 if base == "99205" else 69)
    required_time = first_threshold + max(units - 1, 0) * 15
    time_met = isinstance(total_time, (int, float)) and total_time >= required_time and separate
    if time_met:
        status, deficiencies = "SUPPORTED", []
    elif isinstance(total_time, (int, float)) and total_time >= required_time and not separate:
        status, deficiencies = "BORDERLINE", ["The total time reaches the threshold, but prolonged E/M time is not clearly separated from other reported services."]
    else:
        status, deficiencies = "NOT SUPPORTED", [f"{required_time} minutes of clearly attributable total E/M time were required for {units} unit(s); documented time was {total_time if total_time is not None else 'not stated'}." ]
    return {"status": status, "reasons": [f"Base code: {base}; required total E/M time: {required_time}; documented: {total_time if total_time is not None else 'not stated'}."], "deficiencies": deficiencies}


def _checklist_result(items: list[tuple[str, bool, bool]]) -> dict:
    """items: (label, present, essential)."""
    missing_essential = [label for label, present, essential in items if essential and not present]
    missing_other = [label for label, present, essential in items if not essential and not present]
    present_count = sum(1 for _, present, _ in items if present)
    if missing_essential:
        status = "NOT SUPPORTED"
    elif missing_other:
        status = "BORDERLINE"
    else:
        status = "SUPPORTED"
    reasons = [f"{present_count} of {len(items)} checklist elements were documented."]
    deficiencies = missing_essential + missing_other
    return {"status": status, "reasons": reasons, "deficiencies": deficiencies}


def evaluate_tms(code: str, facts: dict) -> dict:
    common = [
        ("Active diagnosis was not documented.", yes(facts.get("diagnosis_documented")), True),
        ("Medical necessity for TMS was not documented.", yes(facts.get("medical_necessity_documented")), True),
        ("Patient tolerance was not documented.", yes(facts.get("tms_tolerance_documented")), True),
        ("Adverse effects or absence of adverse effects was not documented.", yes(facts.get("tms_adverse_effects_documented")), False),
        ("The treatment plan or continuation plan was not documented.", yes(facts.get("tms_plan_documented")), False),
    ]
    if code == "90867":
        items = common + [
            ("Motor-threshold determination was not documented.", yes(facts.get("tms_motor_threshold_documented")), True),
            ("The motor-threshold method was not documented.", yes(facts.get("tms_motor_threshold_method_documented")), False),
            ("Coil placement or motor hotspot was not documented.", yes(facts.get("tms_coil_placement_documented")), True),
            ("Treatment intensity was not documented.", yes(facts.get("tms_intensity_documented")), True),
            ("The selected protocol or pulse parameters were not documented.", yes(facts.get("tms_pulses_or_protocol_documented")), True),
            ("Baseline symptom severity was not documented.", yes(facts.get("baseline_symptom_scale_documented")), False),
            ("Informed consent was not documented.", yes(facts.get("informed_consent_documented")), False),
        ]
    elif code == "90868":
        items = common + [
            ("Session number was not documented.", yes(facts.get("tms_session_number_documented")), False),
            ("Current coil placement was not documented.", yes(facts.get("tms_coil_placement_documented")), True),
            ("Current treatment intensity was not documented.", yes(facts.get("tms_intensity_documented")), True),
            ("Pulses or treatment protocol were not documented.", yes(facts.get("tms_pulses_or_protocol_documented")), True),
            ("Treatment delivery/completion was not documented.", yes(facts.get("tms_treatment_delivered_documented")), True),
            ("A relevant safety assessment was not documented.", yes(facts.get("tms_safety_assessment_documented")), False),
        ]
    else:
        items = common + [
            ("The clinical reason for repeat motor-threshold determination was not documented.", yes(facts.get("tms_repeat_mt_reason_documented")), True),
            ("Repeat motor-threshold determination was not documented.", yes(facts.get("tms_motor_threshold_documented")), True),
            ("The prior and new thresholds were not compared.", yes(facts.get("tms_prior_and_new_mt_compared")), False),
            ("Updated treatment intensity was not documented.", yes(facts.get("tms_intensity_documented")), True),
            ("Updated parameters or continuation plan were not documented.", yes(facts.get("tms_plan_documented")), True),
        ]
    return _checklist_result(items)


def evaluate_spravato(code: str, facts: dict) -> dict:
    dose = facts.get("spravato_dose_mg")
    expected = 56 if code == "G2082" else 84
    dose_match = isinstance(dose, (int, float)) and int(dose) == expected
    observation = facts.get("spravato_observation_minutes")
    observation_met = isinstance(observation, (int, float)) and observation >= 120
    items = [
        ("The Spravato indication or continued medical necessity was not documented.", yes(facts.get("spravato_indication_documented")) or yes(facts.get("medical_necessity_documented")), True),
        (f"The documented dose did not establish {expected} mg for {code}.", dose_match, True),
        ("Self-administration under healthcare-provider supervision was not documented.", yes(facts.get("spravato_self_administered_under_supervision")), True),
        ("Pre-dose blood pressure was not documented.", yes(facts.get("spravato_pre_dose_blood_pressure")), True),
        ("At least 120 minutes of observation was not explicitly documented.", observation_met, True),
        ("Periodic vital-sign monitoring was not documented.", yes(facts.get("spravato_periodic_vitals_documented")), False),
        ("Sedation assessment was not documented.", yes(facts.get("spravato_sedation_assessment_documented")), False),
        ("Dissociation assessment was not documented.", yes(facts.get("spravato_dissociation_assessment_documented")), False),
        ("Respiratory monitoring or assessment was not documented.", yes(facts.get("spravato_respiratory_monitoring_documented")), False),
        ("Clinical discharge status was not documented.", yes(facts.get("spravato_discharge_status_documented")), True),
        ("Transportation/no-driving instructions were not documented.", yes(facts.get("spravato_transportation_documented")), False),
        ("Follow-up plan was not documented.", yes(facts.get("spravato_followup_plan_documented")), False),
    ]
    result = _checklist_result(items)
    result["reasons"].append(f"Documented dose: {dose if dose is not None else 'not stated'} mg; observation: {observation if observation is not None else 'not stated'} minutes.")
    return result


def evaluate_codes(codes: list[dict], facts: dict, payer: str) -> dict[str, dict]:
    findings = {}
    for entry in codes:
        code, units = entry["code"], entry["units"]
        if code in E_M_RULES:
            finding = evaluate_em(code, facts)
        elif code in PSYCHOTHERAPY_RULES:
            finding = evaluate_psychotherapy(code, facts)
        elif code == "G2211":
            finding = evaluate_g2211(facts, payer, codes)
        elif code in {"99417", "G2212"}:
            finding = evaluate_prolonged(code, units, facts, payer, codes)
        elif code in TMS_CODES:
            finding = evaluate_tms(code, facts)
        elif code in SPRAVATO_CODES:
            finding = evaluate_spravato(code, facts)
        else:
            finding = {"status": "BORDERLINE", "reasons": ["This code is not yet in the deterministic rules engine."], "deficiencies": ["The result requires manual review against the governing excerpts."]}
        finding["units"] = units
        findings[code] = finding
    return findings


def overall_result(findings: dict[str, dict]) -> tuple[str, str, int]:
    statuses = [item["status"] for item in findings.values()]
    if "NOT SUPPORTED" in statuses:
        result = "CORRECTION REQUIRED"
        risk = "HIGH" if statuses.count("NOT SUPPORTED") > 1 else "MEDIUM"
    elif "BORDERLINE" in statuses:
        result, risk = "REVIEW RECOMMENDED", "MEDIUM"
    else:
        result, risk = "PASS", "LOW"
    confidence = max(70, min(98, 96 - statuses.count("BORDERLINE") * 8 - statuses.count("NOT SUPPORTED") * 5))
    return result, risk, confidence


def render_report(codes: list[dict], payer: str, findings: dict[str, dict], explanations: dict, sources: list[str]) -> str:
    result, risk, confidence = overall_result(findings)
    lines = ["# TRIMERA DOCUMENTATION QA", "", "## Provider Intended Billing", ", ".join(entry["raw"] for entry in codes), "", "## Payer", payer, "", "## Overall Result", f"**{result}**", "", "## Documentation Confidence", f"**{confidence}%**", "", "## Overall Audit Risk", f"**{risk}**", "", "## Governing Authority Used", ", ".join(f"[{source}]" for source in sources) or "No authority loaded", "", "---", ""]
    explanation_map = explanations.get("code_explanations", {})
    for entry in codes:
        code, finding = entry["code"], findings[entry["code"]]
        explanation = explanation_map.get(code, {})
        support = explanation.get("support") or " ".join(finding["reasons"])
        deficiencies = explanation.get("deficiencies") or finding["deficiencies"]
        lines.extend([f"## {code}" + (f" × {entry['units']}" if entry["units"] > 1 else ""), f"### {finding['status']}", "", "**Support**", support or "No supporting element was established.", ""])
        if finding["status"] != "SUPPORTED":
            lines.append("**Documentation Deficiencies**")
            lines.extend(f"- {item}" for item in deficiencies)
            lines.append("")
        code_risk = "LOW" if finding["status"] == "SUPPORTED" else "MEDIUM" if finding["status"] == "BORDERLINE" else "HIGH"
        lines.extend(["**Audit Risk**", code_risk, "", "---", ""])
    lines.extend(["## Documentation Quality Improvements", ""])
    improvements = explanations.get("quality_improvements", [])[:3]
    lines.extend([f"- {item}" for item in improvements] or ["None identified."])
    lines.extend(["", "---", "", "## Final Assessment", explanations.get("final_assessment", "The fixed code-level findings above control the result.")])
    return "\n".join(lines)


def reset_qa_session() -> None:
    for key in ["qa_result", "qa_note_text", "qa_codes", "qa_payer", "qa_excerpts", "qa_facts", "qa_findings", "qa_sources", "qa_followup_messages"]:
        st.session_state.pop(key, None)


apply_shared_theme()
require_auth(APP_TITLE, "Internal Trimera Health tool")
render_topbar()

with st.sidebar:
    sidebar_label("Quick actions")
    if st.button("Start new QA review", use_container_width=True):
        reset_qa_session()
        st.rerun()
    if st.button("Sign out", use_container_width=True):
        logout_user()
    sidebar_model(MODEL)
    sidebar_reminder("Reminder", "Code-level outcomes are determined by the fixed rules engine.")

shared_page_header(
    "documentation",
    "Documentation QA",
    "Grounded review using payer, AMA, CMS, and Trimera authority documents.",
)

# Reference files are intentionally not opened here. Loading them at page
# startup made login appear frozen. They are loaded lazily after the user
# clicks Run documentation QA.

payer = st.selectbox("Payer", ["Not specified", "Medicare", "UnitedHealthcare Community Plan / Medicaid", "UnitedHealthcare / Optum Commercial", "BCBS", "Aetna", "Cigna", "Humana", "Other"])
codes_raw = st.text_input("Intended billing", placeholder="99214, 90833, G2211  OR  99215, 99417 x5")
method = st.radio("Clinical note", ["Paste text", "Upload PDF"], horizontal=True)
note_text = ""
if method == "Paste text":
    note_text = st.text_area("Paste completed note", height=360)
else:
    uploaded = st.file_uploader("Upload completed note PDF", type=["pdf"])
    if uploaded:
        try:
            note_text = extract_pdf(uploaded)
            st.success("PDF text extracted.")
        except Exception as exc:
            st.error(f"Could not read PDF: {exc}")

if st.button("Run documentation QA", type="primary", use_container_width=True):
    codes = parse_codes(codes_raw)
    if not codes:
        st.error("Enter at least one intended code.")
        st.stop()
    if not note_text.strip():
        st.error("Paste or upload a note.")
        st.stop()
    if not OPENAI_API_KEY:
        st.error("OPENAI_API_KEY is not configured.")
        st.stop()
    with st.spinner("Loading the relevant authority documents..."):
        try:
            reference_library = load_reference_library(
                payer,
                tuple(entry["code"] for entry in codes),
            )
        except Exception as exc:
            st.error(f"Could not load the reference library: {exc}")
            st.stop()

    excerpts, used_sources = governing_excerpts(
        reference_library,
        payer,
        codes,
    )
    if not excerpts:
        st.error("No governing reference excerpts were found. Confirm the files exist in the reference folder with the expected filenames.")
        st.stop()
    client = OpenAI(api_key=OPENAI_API_KEY)
    extraction_input = f"INTENDED BILLING:\n{chr(10).join(entry['raw'] for entry in codes)}\n\nPAYER:\n{payer}\n\nCOMPLETED CLINICAL NOTE:\n{note_text}"
    with st.spinner("Extracting documented facts..."):
        try:
            extraction_response = client.responses.create(model=MODEL, instructions=FACT_EXTRACTION_PROMPT, input=extraction_input)
            facts = clean_json(extraction_response.output_text)
        except Exception as exc:
            st.error(f"Fact extraction failed: {exc}")
            st.stop()
    findings = evaluate_codes(codes, facts, payer)
    explanation_input = f"PAYER:\n{payer}\n\nINTENDED BILLING:\n{json.dumps(codes, indent=2)}\n\nFIXED CODE FINDINGS:\n{json.dumps(findings, indent=2)}\n\nEXTRACTED FACTS:\n{json.dumps(facts, indent=2)}\n\nGOVERNING EXCERPTS:\n{excerpts}\n\nCOMPLETED NOTE:\n{note_text}"
    with st.spinner("Writing grounded feedback..."):
        try:
            explanation_response = client.responses.create(
                model=MODEL,
                instructions=with_web_research(EXPLANATION_PROMPT),
                input=explanation_input,
                tools=WEB_SEARCH_TOOLS,
            )
            explanations = clean_json(explanation_response.output_text)
        except Exception as exc:
            explanations = {"code_explanations": {}, "quality_improvements": [], "final_assessment": f"The fixed rules engine completed the review, but the explanation layer failed: {exc}"}
    report = render_report(codes, payer, findings, explanations, used_sources)
    st.session_state.update({"qa_result": report, "qa_note_text": note_text, "qa_codes": codes, "qa_payer": payer, "qa_excerpts": excerpts, "qa_facts": facts, "qa_findings": findings, "qa_sources": used_sources, "qa_followup_messages": []})

if st.session_state.get("qa_result"):
    report = st.session_state["qa_result"]
    st.divider()
    st.subheader("QA result")
    st.markdown(report)
    st.download_button("Download report", data=report, file_name="trimera_documentation_qa.md", mime="text/markdown", use_container_width=True)
    with st.expander("Technical fact extraction"):
        st.json(st.session_state["qa_facts"])
    st.divider()
    st.subheader("💬 Ask Trimera about this QA review")
    st.caption("Ask why a fixed result was reached, what documentation is missing, or request concise provider education. Trimera automatically checks reputable current web sources when they add relevant context and cites web-derived information.")
    for message in st.session_state.get("qa_followup_messages", []):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    followup_question = st.chat_input("Ask a follow-up about this documentation review...")
    if followup_question:
        st.session_state["qa_followup_messages"].append({"role": "user", "content": followup_question})
        with st.chat_message("user"):
            st.markdown(followup_question)
        conversation = "\n\n".join(f"{message['role'].upper()}:\n{message['content']}" for message in st.session_state["qa_followup_messages"])
        followup_context = f"PAYER:\n{st.session_state['qa_payer']}\n\nINTENDED BILLING:\n{json.dumps(st.session_state['qa_codes'], indent=2)}\n\nFIXED FINDINGS:\n{json.dumps(st.session_state['qa_findings'], indent=2)}\n\nEXTRACTED FACTS:\n{json.dumps(st.session_state['qa_facts'], indent=2)}\n\nORIGINAL REPORT:\n{st.session_state['qa_result']}\n\nGOVERNING EXCERPTS:\n{st.session_state['qa_excerpts']}\n\nCOMPLETED NOTE:\n{st.session_state['qa_note_text']}\n\nFOLLOW-UP CONVERSATION:\n{conversation}"
        client = OpenAI(api_key=OPENAI_API_KEY)
        with st.chat_message("assistant"):
            with st.spinner("Reviewing the fixed QA result..."):
                try:
                    response = client.responses.create(model=MODEL, instructions=with_web_research(FOLLOWUP_PROMPT), input=followup_context, tools=WEB_SEARCH_TOOLS)
                    answer = response.output_text
                    st.markdown(answer)
                except Exception as exc:
                    st.error(f"OpenAI request failed: {exc}")
                    st.stop()
        st.session_state["qa_followup_messages"].append({"role": "assistant", "content": answer})
